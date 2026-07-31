"""Build the printable DOCX worksheets used by the local human-task pilot."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).parents[1]
OUTPUT_DIR = ROOT / "portal_materials"
BLUE = "2E74B5"
NAVY = "0B2545"
MUTED = "5C6F7A"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
RED = "9B1C1C"

DOCUMENTS = [
    {
        "task_id": "DEF-002",
        "filename": "DEF-002_kanonikus_auditjelentes_review.docx",
        "title": "Kanonikus auditjelentés – emberi review",
        "purpose": (
            "Annak írásos megerősítése, hogy az aláírt auditjelentés és a "
            "D-025 döntési rekord ugyanarra a használható, kanonikus példányra vonatkozik."
        ),
        "owner": "Lángi Zoltán",
        "reviewer": "Lángi Zoltán vagy kijelölt reviewer",
        "source_refs": "SRC-008; D-025",
        "steps": [
            "Nyissa meg az aláírt auditjelentést és a D-025 döntési rekordot.",
            "Hasonlítsa össze a fájl nevét, dátumát, tartalmi azonosítóit és a dokumentált hash-t.",
            "Rögzítse az eredményt, majd írja alá ezt a review-lapot.",
            "A beszkennelt vagy elektronikusan aláírt végleges példányt töltse fel a védett NIS2 SharePoint-tárba.",
        ],
        "checks": [
            "A jelentés azonosítója és dátuma egyezik",
            "A D-025 ugyanarra a példányra hivatkozik",
            "Az eltérések vagy bizonytalanságok le vannak írva",
            "A kanonikus példány elfogadható / javításra visszaadandó",
        ],
    },
    {
        "task_id": "DEF-004",
        "filename": "DEF-004_auditigazolas_olvashatosagi_ellenorzes.docx",
        "title": "Auditigazolás – olvashatósági ellenőrzés",
        "purpose": (
            "Annak dokumentálása, hogy az SRC-008 dokumentum 388. oldala "
            "emberi PDF-megjelenítőben teljes egészében olvasható."
        ),
        "owner": "Pásztor András",
        "reviewer": "Lángi Zoltán vagy kijelölt reviewer",
        "source_refs": "SRC-008:p388",
        "steps": [
            "Nyissa meg az SRC-008 dokumentumot egy jóváhagyott PDF-megjelenítőben.",
            "Lépjen a 388. oldalra, és ellenőrizze a szövegek, dátumok, aláírások és bélyegzők olvashatóságát.",
            "Jegyezze fel a használt programot és verziót, valamint minden észlelt hibát.",
            "Írja alá a lapot, majd a végleges példányt töltse fel a védett NIS2 SharePoint-tárba.",
        ],
        "checks": [
            "Az oldal teljes terjedelmében megjelenik",
            "A szöveg és a dátumok olvashatók",
            "Az aláírások és bélyegzők felismerhetők",
            "Nincs levágás, hibás karakter vagy üres tartalmi rész",
        ],
    },
    {
        "task_id": "DEF-005",
        "filename": "DEF-005_RACI_es_IBF_kijelolesi_csomag.docx",
        "title": "RACI- és IBF-kijelölési csomag",
        "purpose": (
            "A jóváhagyott NIS2-szerepkiosztás formális rögzítése, "
            "a hatáskörök, helyettesítések és aláírások ellenőrzésével."
        ),
        "owner": "Kóczán Mónika",
        "reviewer": "Lángi Zoltán",
        "source_refs": "D-027; HUMAN_EXECUTION_PACKAGE",
        "steps": [
            "Ellenőrizze a szerepkiosztásban szereplő személyeket és szervezeteket.",
            "Rögzítse a döntési, végrehajtási, véleményezési és tájékoztatási felelősségeket.",
            "Ellenőrizze a helyettesítést, a partneri szerepek korlátait és a vezetői jóváhagyást.",
            "Írassa alá a kijelölési csomagot, majd töltse fel a védett NIS2 SharePoint-tárba.",
        ],
        "checks": [
            "IBF: Lángi Zoltán",
            "Jogi reviewer: Dr. Berta Brigitta",
            "Projektvezető: Kóczán Mónika",
            "Infrastruktúra és incidenskezelés: belső kontrollgazda + Serversystem Kft.",
            "HR: Koncz Erika",
            "Fizikai védelem: Német Péter",
        ],
    },
    {
        "task_id": "DEF-006",
        "filename": "DEF-006_IBF_alkalmassagi_jogi_review.docx",
        "title": "IBF-alkalmasság – jogi review",
        "purpose": (
            "Az IBF-re alkalmazandó követelmények és azok teljesülésének jogi "
            "ellenőrzése úgy, hogy személyes okirat ne kerüljön a Git-re vagy a portál munkanaplójába."
        ),
        "owner": "Dr. Berta Brigitta",
        "reviewer": "Lángi Zoltán",
        "source_refs": "DEF-006; SECURITY_BOUNDARIES",
        "steps": [
            "Határozza meg a metALCOM Zrt.-re és az IBF-re alkalmazandó jogcímet.",
            "Ellenőrizze a végzettségi, tapasztalati, összeférhetetlenségi és továbbképzési követelményeket.",
            "A személyes okiratokat csak védett tárban kezelje; ezen a lapon kizárólag az ellenőrzés eredményét és a tárhivatkozást rögzítse.",
            "Írja alá a jogi review-t, majd töltse fel a védett NIS2 SharePoint-tárba.",
        ],
        "checks": [
            "Az alkalmazandó jogcím dokumentált",
            "A szükséges követelmények tételesen ellenőrizve",
            "Az eltérés vagy pótlási igény dokumentált",
            "Személyes okirat nem került Gitbe vagy nyilvános helyre",
        ],
    },
    {
        "task_id": "DEF-007",
        "filename": "DEF-007_kontrollgazda_es_partner_review.docx",
        "title": "Kontrollgazda-kijelölés és partneri review",
        "purpose": (
            "A belső infrastruktúra- és incidenskezelési kontrollgazda kijelölése, "
            "valamint a Serversystem Kft. releváns szerződéses kötelezettségeinek áttekintése."
        ),
        "owner": "Kóczán Mónika",
        "reviewer": "Dr. Berta Brigitta és Lángi Zoltán",
        "source_refs": "D-027; DEF-007",
        "steps": [
            "Jelölje ki a metALCOM belső kontrollgazdáját és helyettesét.",
            "Tekintsék át a Serversystem Kft. incidens-, napló-, hozzáférési, értesítési és auditkötelezettségeit.",
            "Rögzítsék a hiányzó vagy pontosítandó szerződéses pontokat és a felelőst.",
            "Írják alá a kijelölést és a review-t, majd töltsék fel a védett NIS2 SharePoint-tárba.",
        ],
        "checks": [
            "Belső kontrollgazda és helyettes kijelölve",
            "Incidensbejelentési kötelezettség ellenőrizve",
            "Napló- és hozzáféréskezelési kötelezettség ellenőrizve",
            "Auditjog és bizonyíték-átadás ellenőrizve",
            "Hiányokhoz felelős és céldátum rendelve",
        ],
    },
]


def set_cell_fill(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_width(cell, dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def configure_table(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_mar = cell._tc.get_or_add_tcPr().find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                cell._tc.get_or_add_tcPr().append(tc_mar)
            for edge, value in (
                ("top", 80),
                ("bottom", 80),
                ("start", 120),
                ("end", 120),
            ):
                node = OxmlElement(f"w:{edge}")
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")
                tc_mar.append(node)


def style_run(run, *, size: float = 11, bold: bool = False, color: str = "000000") -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_field_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    configure_table(table, [2700, 6660])
    for index, (label, value) in enumerate(rows):
        set_cell_fill(table.cell(index, 0), LIGHT_BLUE)
        style_run(table.cell(index, 0).paragraphs[0].add_run(label), bold=True, color=NAVY)
        style_run(table.cell(index, 1).paragraphs[0].add_run(value))


def add_lines(doc: Document, count: int = 3) -> None:
    for _ in range(count):
        paragraph = doc.add_paragraph("________________________________________________________________________________")
        paragraph.paragraph_format.space_after = Pt(4)
        style_run(paragraph.runs[0], size=9, color=MUTED)


def add_footer(section, task_id: str) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_run(
        paragraph.add_run(
            f"{task_id} · Emberi munkalap · TERVEZET – kitöltés és jóváhagyás előtt nem evidencia"
        ),
        size=8,
        color=MUTED,
    )


def build_document(spec: dict[str, object]) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, BLUE),
        ("Heading 2", 13, 14, 7, BLUE),
        ("Heading 3", 12, 10, 5, NAVY),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    for name in ("List Number", "List Bullet"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    style_run(header.add_run("metALCOM NIS2 · emberi végrehajtási csomag"), size=9, bold=True, color=MUTED)
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_footer(section, str(spec["task_id"]))

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(2)
    style_run(kicker.add_run(str(spec["task_id"])), size=10, bold=True, color=BLUE)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    style_run(title.add_run(str(spec["title"])), size=23, bold=True, color=NAVY)
    warning = doc.add_paragraph()
    warning.paragraph_format.space_after = Pt(12)
    style_run(
        warning.add_run("TERVEZET · kitöltés és aláírás előtt nem evidencia"),
        size=9,
        bold=True,
        color=RED,
    )

    add_field_table(
        doc,
        [
            ("Feladat", str(spec["task_id"])),
            ("Felelős", str(spec["owner"])),
            ("Reviewer", str(spec["reviewer"])),
            ("Forrás", str(spec["source_refs"])),
            ("Kitöltés dátuma", "________________________________"),
        ],
    )
    doc.add_heading("Mire való ez a lap?", level=2)
    doc.add_paragraph(str(spec["purpose"]))
    doc.add_heading("Elvégzendő lépések", level=2)
    for step in spec["steps"]:
        doc.add_paragraph(str(step), style="List Number")

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_footer(doc.sections[-1], str(spec["task_id"]))
    doc.add_heading("Ellenőrzési pontok", level=1)
    checks = list(spec["checks"])
    table = doc.add_table(rows=len(checks) + 1, cols=3)
    table.style = "Table Grid"
    configure_table(table, [6120, 1620, 1620])
    for index, value in enumerate(("Ellenőrzési pont", "Megfelel", "Nem felel meg")):
        set_cell_fill(table.cell(0, index), LIGHT_BLUE)
        style_run(table.cell(0, index).paragraphs[0].add_run(value), bold=True, color=NAVY)
    for row_index, check in enumerate(checks, 1):
        style_run(table.cell(row_index, 0).paragraphs[0].add_run(str(check)))
        style_run(table.cell(row_index, 1).paragraphs[0].add_run("Igen / N.a."))
        style_run(table.cell(row_index, 2).paragraphs[0].add_run("Nem / N.a."))

    doc.add_heading("Megállapítás és szükséges intézkedés", level=2)
    add_lines(doc, 3)
    doc.add_heading("Döntés", level=2)
    decision = doc.add_table(rows=2, cols=2)
    decision.style = "Table Grid"
    configure_table(decision, [4680, 4680])
    style_run(decision.cell(0, 0).paragraphs[0].add_run("Elfogadható / előterjeszthető"), bold=True)
    style_run(decision.cell(0, 1).paragraphs[0].add_run("Javításra vagy pótlásra visszaadva"), bold=True)
    style_run(decision.cell(1, 0).paragraphs[0].add_run("Jelölés: __________________"))
    style_run(decision.cell(1, 1).paragraphs[0].add_run("Jelölés: __________________"))

    doc.add_heading("Végleges dokumentum nyilvántartása", level=2)
    add_field_table(
        doc,
        [
            ("Végleges fájlnév", ""),
            ("SharePoint-hivatkozás", ""),
            ("SHA-256", ""),
            ("Feltöltő / dátum", ""),
        ],
    )
    doc.add_paragraph(
        "A helyi portálon csatolt fájl csak előkészített munkapéldány. "
        "A formális evidenciához a végleges fájlt a védett NIS2 SharePoint-tárba "
        "kell feltölteni, majd a tényleges fájlhivatkozást és a SHA-256 értéket kell rögzíteni."
    )
    doc.add_heading("Aláírások", level=2)
    signatures = doc.add_table(rows=3, cols=2)
    signatures.style = "Table Grid"
    configure_table(signatures, [4680, 4680])
    for col, label in enumerate(("Előkészítő / felelős", "Reviewer / jóváhagyó")):
        set_cell_fill(signatures.cell(0, col), LIGHT_GRAY)
        style_run(signatures.cell(0, col).paragraphs[0].add_run(label), bold=True, color=NAVY)
        style_run(signatures.cell(1, col).paragraphs[0].add_run("Név: ______________________________"))
        style_run(signatures.cell(2, col).paragraphs[0].add_run("Dátum, aláírás: ____________________"))

    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.comments = ""
    doc.core_properties.keywords = ""
    doc.core_properties.title = str(spec["title"])
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output = OUTPUT_DIR / str(spec["filename"])
    doc.save(output)
    return output


def main() -> None:
    outputs = [build_document(spec) for spec in DOCUMENTS]
    for output in outputs:
        print(output)


if __name__ == "__main__":
    main()
