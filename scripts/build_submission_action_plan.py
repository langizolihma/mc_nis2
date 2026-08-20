"""Build the submission-ready NIS2 action plan DOCX from canonical repository data."""

from __future__ import annotations

import csv
import hashlib
import json
from collections import Counter, defaultdict
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DATA = ROOT / "data"
OUT = ROOT / "submission_materials"
DOCX_PATH = OUT / "NIS2_cselekvesi_terv_beadasra_elokeszitett_2026-08-19.docx"
AS_OF = date(2026, 8, 19)

NAVY = "17365D"
BLUE = "1F4E78"
TEAL = "0F6B78"
LIGHT_BLUE = "DDEBF7"
PALE_BLUE = "EAF3F8"
RED = "C00000"
PALE_RED = "FCE4D6"
GOLD = "BF9000"
PALE_GOLD = "FFF2CC"
GREEN = "548235"
PALE_GREEN = "E2F0D9"
GRAY = "666666"
LIGHT_GRAY = "F2F2F2"
WHITE = "FFFFFF"

FIXED_DATE_REQUIRED = {
    "A-008", "A-022", "A-023", "A-024", "A-025", "A-026", "A-027", "A-028", "A-042"
} | {f"A-{index:03d}" for index in range(43, 128)}
UNVERIFIED_SOURCE_ACTIONS = {"A-022", "A-024", "A-026", "A-027", "A-028"}
SUBMISSION_BLOCKERS = {
    "A-004": "BEADÁSI BLOKKOLÓ – a gépi finding-regiszter G1 mintavétele és reviewer sign-offja hiányzik.",
    "A-005": "BEADÁSI BLOKKOLÓ – a 85 új, pontos kontrollszintű intézkedési javaslat és mapping G1 owner review-ja hiányzik.",
    "A-036": "BEADÁSI BLOKKOLÓ – az aláírt RACI, a formális szerepkijelölések és a kapcsolódó evidencia hiányzik.",
    "A-006": "BEADÁSI BLOKKOLÓ – a végleges terv G1/G2/G4 felülvizsgálata és jóváhagyása még nem történt meg.",
}


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def load_actions() -> list[dict[str, str]]:
    with (DATA / "actions.csv").open(encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))
    return sorted(rows, key=lambda row: int(row["action_id"].split("-")[1]))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 70, start: int = 85, bottom: int = 70, end: int = 85) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_width(cell, inches: float) -> None:
    cell.width = Inches(inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_fixed(table) -> None:
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Frissítendő mező"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, text, end):
        run._r.append(element)


def set_repeat_header_distance(section) -> None:
    section.header_distance = Cm(0.6)
    section.footer_distance = Cm(0.6)


def configure_section(section, landscape: bool = False) -> None:
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        section.left_margin = Inches(0.62)
        section.right_margin = Inches(0.62)
        section.top_margin = Inches(0.72)
        section.bottom_margin = Inches(0.72)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        section.top_margin = Inches(0.82)
        section.bottom_margin = Inches(0.82)
    set_repeat_header_distance(section)


def add_header_footer(section, compact: bool = False) -> None:
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("metALCOM Zrt.  |  NIS2 cselekvési terv")
    run.font.name = "Aptos"
    run.font.size = Pt(8 if compact else 8.5)
    run.font.color.rgb = RGBColor.from_string(GRAY)
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(f"BELSŐ MEGFELELŐSÉGI DOKUMENTUM  •  {AS_OF.isoformat()}  •  ")
    run.font.name = "Aptos"
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor.from_string(GRAY)
    add_field(paragraph, "PAGE")
    paragraph.add_run(" / ")
    add_field(paragraph, "NUMPAGES")


def style_document(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string("222222")
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08
    for name, size, color, before, after in (
        ("Title", 27, NAVY, 0, 10),
        ("Subtitle", 13, TEAL, 0, 8),
        ("Heading 1", 18, NAVY, 16, 7),
        ("Heading 2", 13, BLUE, 12, 5),
        ("Heading 3", 10.5, TEAL, 8, 3),
    ):
        style = styles[name]
        style.font.name = "Aptos Display" if name != "Heading 3" else "Aptos"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    styles["Title"].paragraph_format.space_before = Pt(0)


def add_rule(paragraph, color: str = TEAL, size: int = 18) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def add_banner(doc: Document, text: str, fill: str, color: str = "222222") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 120, 150, 120, 150)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(color)


def add_key_value_table(doc: Document, rows: list[tuple[str, str]], widths=(1.8, 4.7)) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_fixed(table)
    for key, value in rows:
        cells = table.add_row().cells
        set_row_cant_split(table.rows[-1])
        set_cell_width(cells[0], widths[0])
        set_cell_width(cells[1], widths[1])
        set_cell_shading(cells[0], LIGHT_BLUE)
        for cell in cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(key)
        r.bold = True
        r.font.size = Pt(8.5)
        p = cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        r.font.size = Pt(8.5)


def add_bullet(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)


def pending_notes(action: dict[str, str], action_by_id: dict[str, dict[str, str]]) -> list[tuple[str, str]]:
    action_id = action["action_id"]
    notes: list[tuple[str, str]] = []
    if action_id in SUBMISSION_BLOCKERS:
        notes.append(("red", SUBMISSION_BLOCKERS[action_id]))
    target = action.get("target_date", "").strip()
    if target and action["status"] not in {"DONE", "CLOSED", "CANCELLED"}:
        if date.fromisoformat(target) < AS_OF:
            notes.append(("red", f"LEJÁRT BELSŐ CÉLDÁTUM – az eredeti {target} dátumhoz státusz- vagy újraütemezési döntés kell."))
    if action_id in FIXED_DATE_REQUIRED and not target:
        notes.append(("gold", "FÜGGŐ – a relatív/eseményalapú határidőhöz konkrét dátum vagy írásban jóváhagyott indoklás szükséges."))
    if action_id in UNVERIFIED_SOURCE_ACTIONS or action.get("source_confidence") == "unverified_internal":
        notes.append(("gold", "FÜGGŐ – az SRC-004 belső állítás read-only validációig csak feltételes információ, nem auditált tény."))
    dependencies = [item.strip() for item in action.get("dependencies", "").split(";") if item.strip()]
    incomplete = [item for item in dependencies if action_by_id.get(item, {}).get("status") != "DONE"]
    if incomplete:
        notes.append(("gold", "FÜGGŐ ELŐFELTÉTEL – " + ", ".join(incomplete) + "."))
    if action.get("status") != "DONE" and not notes:
        notes.append(("gold", "NYITOTT VÉGREHAJTÁS – a feladat és evidenciája még nem lezárt."))
    if action.get("status") == "DONE":
        notes.append(("green", "LEZÁRT NYILVÁNTARTÁSI TÉTEL – az elfogadott döntési és evidenciarekord szerint."))
    return notes


def add_action_cell_text(cell, action: dict[str, str], notes: list[tuple[str, str]], column: int) -> None:
    cell.text = ""
    if column == 0:
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(action["action_id"])
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string(NAVY)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(action["priority"])
        r.bold = True
        r.font.size = Pt(7)
    elif column == 1:
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(action["workstream"])
        r.bold = True
        r.font.size = Pt(8.5)
        p = cell.add_paragraph(action["task"])
        p.paragraph_format.space_after = Pt(3)
        for run in p.runs:
            run.font.size = Pt(8)
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("Eredmény: ")
        r.bold = True
        r.font.size = Pt(7.5)
        r = p.add_run(action["deliverable"])
        r.font.size = Pt(7.5)
        r.font.color.rgb = RGBColor.from_string(GRAY)
    elif column == 2:
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run("Felelős\n")
        r.bold = True
        r.font.size = Pt(7)
        r.font.color.rgb = RGBColor.from_string(GRAY)
        r = p.add_run(action["human_owner"])
        r.font.size = Pt(8)
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("Jóváhagyó\n")
        r.bold = True
        r.font.size = Pt(7)
        r.font.color.rgb = RGBColor.from_string(GRAY)
        r = p.add_run(action["human_approver"])
        r.font.size = Pt(8)
    elif column == 3:
        target = action.get("target_date") or "nincs rögzítve"
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(target)
        r.bold = True
        r.font.size = Pt(8)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(action["status"])
        r.bold = True
        r.font.size = Pt(7.5)
        deadline_basis = action["deadline_basis"]
        if deadline_basis == "D-035_repeat_audit_minus_60_complexity_schedule":
            deadline_basis = "D-035 ütemezési baseline"
        p = cell.add_paragraph(deadline_basis)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            run.font.size = Pt(6.8)
            run.font.color.rgb = RGBColor.from_string(GRAY)
    else:
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run("Elvárt evidencia: ")
        r.bold = True
        r.font.size = Pt(7.5)
        r = p.add_run(action["evidence_required"])
        r.font.size = Pt(7.5)
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run("Forrás: ")
        r.bold = True
        r.font.size = Pt(7)
        r = p.add_run(
            f"{action['source_ref']} ({action['source_confidence']}); "
            f"család: {action['requirement_family']}; kontroll: {action['control_ref'] or '—'}; EIR: {action['scope_eir']}"
        )
        r.font.size = Pt(7)
        r.font.color.rgb = RGBColor.from_string(GRAY)
        for level, text in notes:
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(text)
            r.bold = True
            r.font.size = Pt(7.3)
            if level == "red":
                r.font.color.rgb = RGBColor.from_string(RED)
                r.font.highlight_color = WD_COLOR_INDEX.PINK
            elif level == "gold":
                r.font.color.rgb = RGBColor.from_string(GOLD)
                r.font.highlight_color = WD_COLOR_INDEX.YELLOW
            else:
                r.font.color.rgb = RGBColor.from_string(GREEN)
        if action.get("notes"):
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run("Megjegyzés: ")
            r.bold = True
            r.font.size = Pt(6.8)
            r = p.add_run(action["notes"])
            r.font.size = Pt(6.8)
            r.font.color.rgb = RGBColor.from_string(GRAY)


def add_action_matrix(doc: Document, actions: list[dict[str, str]]) -> None:
    action_by_id = {row["action_id"]: row for row in actions}
    grouped: dict[int, list[dict[str, str]]] = defaultdict(list)
    for row in actions:
        grouped[int(row["requirement_family"].split(";")[0])].append(row)
    chunk_size = 1
    for family in sorted(grouped):
        family_actions = grouped[family]
        for offset in range(0, len(family_actions), chunk_size):
            suffix = " (folytatás)" if offset else ""
            page_marker = doc.add_paragraph()
            page_marker.paragraph_format.page_break_before = True
            page_marker.paragraph_format.space_after = Pt(0)
            page_marker.paragraph_format.line_spacing = Pt(1)
            table = doc.add_table(rows=2, cols=5)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            set_table_fixed(table)
            widths = (0.55, 3.05, 1.18, 1.05, 3.48)
            headers = ("ID", "Feladat és elvárt eredmény", "Felelősség", "Határidő / állapot", "Evidencia, forrás és függő elem")
            segment_actions = family_actions[offset:offset + chunk_size]
            banner = table.rows[0]
            set_row_cant_split(banner)
            banner_cell = banner.cells[0].merge(banner.cells[-1])
            set_cell_shading(banner_cell, LIGHT_BLUE)
            set_cell_margins(banner_cell, 150, 100, 150, 100)
            banner_text = f"{family}. követelménycsalád{suffix}  |  {segment_actions[0]['action_id']}–{segment_actions[-1]['action_id']}"
            banner_paragraph = banner_cell.paragraphs[0]
            banner_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            banner_paragraph.paragraph_format.space_after = Pt(0)
            banner_run = banner_paragraph.add_run(banner_text)
            banner_run.bold = True
            banner_run.font.size = Pt(9)
            banner_run.font.color.rgb = RGBColor.from_string(NAVY)
            header = table.rows[1]
            set_repeat_table_header(header)
            header.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            for index, cell in enumerate(header.cells):
                set_cell_width(cell, widths[index])
                set_cell_shading(cell, NAVY)
                set_cell_margins(cell, 75, 80, 75, 80)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.keep_with_next = True
                r = p.add_run(headers[index])
                r.bold = True
                r.font.size = Pt(7.3)
                r.font.color.rgb = RGBColor.from_string(WHITE)
            for action in segment_actions:
                row = table.add_row()
                set_row_cant_split(row)
                notes = pending_notes(action, action_by_id)
                row_fill = PALE_GREEN if action["status"] == "DONE" else WHITE
                if any(level == "red" for level, _ in notes):
                    row_fill = PALE_RED
                elif any(level == "gold" for level, _ in notes):
                    row_fill = PALE_GOLD
                for index, cell in enumerate(row.cells):
                    set_cell_width(cell, widths[index])
                    set_cell_margins(cell, 60, 70, 60, 70)
                    set_cell_shading(cell, row_fill if index in (0, 3, 4) else WHITE)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                    add_action_cell_text(cell, action, notes, index)
            doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_family_coverage(doc: Document, actions: list[dict[str, str]]) -> None:
    """Add a non-duplicating 19-family cross-reference before the action register."""
    coverage: dict[int, list[str]] = {family: [] for family in range(1, 20)}
    for action in actions:
        for raw_family in action["requirement_family"].split(";"):
            coverage[int(raw_family)].append(action["action_id"])
    doc.add_heading("7.1 Követelménycsalád-lefedettség", level=2)
    p = doc.add_paragraph(
        "A több családhoz kapcsolódó intézkedés a törzsnyilvántartásban egyszer, az elsődleges családjánál jelenik meg; "
        "az alábbi kereszthivatkozás mind a 19 család teljes kapcsolati körét mutatja."
    )
    p.paragraph_format.space_after = Pt(5)
    p.runs[0].font.size = Pt(8.2)
    p.runs[0].font.color.rgb = RGBColor.from_string(GRAY)
    for family in range(1, 20):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.25)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{family}. követelménycsalád: ")
        r.bold = True
        r.font.size = Pt(8.2)
        r.font.color.rgb = RGBColor.from_string(NAVY)
        r = p.add_run(", ".join(coverage[family]))
        r.font.size = Pt(8.2)


def build() -> Path:
    actions = load_actions()
    action_count = len(actions)
    action_by_id = {row["action_id"]: row for row in actions}
    counts = Counter(row["status"] for row in actions)
    overdue = [
        row for row in actions
        if row.get("target_date")
        and row["status"] not in {"DONE", "CLOSED", "CANCELLED"}
        and date.fromisoformat(row["target_date"]) < AS_OF
    ]
    with (DATA / "project_dates.json").open(encoding="utf-8") as handle:
        dates = json.load(handle)
    with (DATA / "source_register.json").open(encoding="utf-8") as handle:
        sources = {item["id"]: item for item in json.load(handle)}
    authority_source = sources["SRC-001"]

    OUT.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    configure_section(doc.sections[0], landscape=False)
    add_header_footer(doc.sections[0])

    # Cover: editorial_cover pattern.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("metALCOM Zrt.")
    r.bold = True
    r.font.name = "Aptos"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(TEAL)
    p = doc.add_paragraph()
    add_rule(p, TEAL, 22)
    p.paragraph_format.space_after = Pt(50)
    p = doc.add_paragraph(style="Title")
    p.add_run("NIS2 helyreállítási\ncselekvési terv")
    p = doc.add_paragraph(style="Subtitle")
    p.add_run("Beadásra előkészített, formailag végleges tervezet")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run(f"{action_count} intézkedés  •  19 követelménycsalád  •  határidő: 2026. szeptember 24.")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(60)
    add_banner(
        doc,
        "FIGYELEM – A piros és sárga jelölésű elemek lezárásáig a dokumentum nem írható alá és nem nyújtható be végleges tervként.",
        PALE_RED,
        RED,
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(24)
    add_key_value_table(doc, [
        ("Dokumentumállapot", "PROPOSAL – G1/G2/G4 emberi jóváhagyásra vár"),
        ("Verzió", "1.0 – beadásra előkészített tervezet"),
        ("Állapot dátuma", AS_OF.isoformat()),
        ("Végrehajtásért felelős", "Pásztor András"),
        ("Jóváhagyó", "Lángi Zoltán"),
        ("Kanonikus beadási határidő", dates["action_plan_deadline"]),
    ])
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Belső megfelelőségi és hatósági beadás-előkészítési dokumentum")
    r.italic = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(GRAY)

    doc.add_page_break()
    doc.add_heading("Dokumentumkontroll", level=1)
    add_key_value_table(doc, [
        ("Forrás-nyilvántartás", f"data/actions.csv – {action_count} intézkedés"),
        ("Forrás SHA-256", sha256(DATA / "actions.csv")),
        ("Dátumbaseline", "data/project_dates.json"),
        ("Dátumbaseline SHA-256", sha256(DATA / "project_dates.json")),
        ("Hatósági döntés", "SRC-001 – " + authority_source["file"]),
        ("Hatósági döntés SHA-256", authority_source["sha256"]),
        ("Védett forráshivatkozás", authority_source["internal_uri"]),
        ("Kézhezvételi baseline", dates["receipt_date"] + " – D-022/D-031 szerint elfogadott"),
        ("Beadási határidő", dates["action_plan_deadline"]),
        ("Következő ismétlő audit célja", "2027-09-30 (D-021); végső korlát: " + dates["repeat_audit_latest"]),
        ("Minősítés", "AI által előállított javaslat; csak emberi felülvizsgálat és aláírás után használható végleges dokumentumként."),
    ])

    doc.add_heading("Jelölések", level=2)
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row, (fill, label, desc) in zip(table.rows, (
        (PALE_RED, "PIROS – BLOKKOLÓ / LEJÁRT", "Beadás előtt emberi döntés, evidencia vagy újraütemezés szükséges."),
        (PALE_GOLD, "SÁRGA – FÜGGŐ", "Nyitott végrehajtás, dátum, előfeltétel vagy feltételes forrás."),
        (PALE_GREEN, "ZÖLD – IGAZOLT", "Elfogadott döntési és evidenciarekord alapján lezárt tétel."),
    )):
        set_row_cant_split(row)
        set_cell_shading(row.cells[0], fill)
        for cell in row.cells:
            set_cell_margins(cell)
        r = row.cells[0].paragraphs[0].add_run(label)
        r.bold = True
        r.font.size = Pt(8.5)
        r = row.cells[1].paragraphs[0].add_run(desc)
        r.font.size = Pt(8.5)

    doc.add_heading("Tartalomjegyzék", level=1)
    toc = doc.add_paragraph()
    add_field(toc, 'TOC \\o "1-2" \\h \\z \\u')
    p = doc.add_paragraph("A tartalomjegyzék Wordben a mezők frissítésével (Ctrl+A, F9) újraszámítható.")
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor.from_string(GRAY)

    doc.add_page_break()
    doc.add_heading("1. Vezetői összefoglaló", level=1)
    p = doc.add_paragraph(
        "A terv célja, hogy a hatósági és auditmegállapításokra 19 követelménycsalád szerint, "
        "név szerinti felelőssel, határidővel, kézzelfogható eredménnyel és ellenőrizhető evidenciával adjon végrehajtható választ. "
        f"A dokumentum a teljes {action_count} tételes nyilvántartást tartalmazza; a hiányzó emberi döntéseket nem rejti el, hanem külön jelöli."
    )
    p.paragraph_format.space_after = Pt(8)
    metrics = doc.add_table(rows=2, cols=4)
    metrics.alignment = WD_TABLE_ALIGNMENT.CENTER
    metrics.style = "Table Grid"
    metric_values = (
        (str(action_count), "INTÉZKEDÉS", PALE_BLUE),
        ("19/19", "KÖVETELMÉNYCSALÁD", PALE_BLUE),
        ("0", "HARD VALIDÁCIÓS HIBA", PALE_GREEN),
        (str(len(FIXED_DATE_REQUIRED)), "DÁTUMJÓVÁHAGYÁS", PALE_GOLD),
    )
    for col, (value, label, fill) in enumerate(metric_values):
        set_cell_shading(metrics.cell(0, col), fill)
        set_cell_shading(metrics.cell(1, col), fill)
        for row in (0, 1):
            set_cell_margins(metrics.cell(row, col), 90, 70, 90, 70)
        p = metrics.cell(0, col).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        r.bold = True
        r.font.size = Pt(18)
        r.font.color.rgb = RGBColor.from_string(NAVY)
        p = metrics.cell(1, col).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(6.8)
        r.font.color.rgb = RGBColor.from_string(GRAY)

    doc.add_heading("Aktuális állapot", level=2)
    add_bullet(doc, f"Lezárt: {counts.get('DONE', 0)}; folyamatban: {counts.get('IN_PROGRESS', 0)}; új: {counts.get('NEW', 0)} intézkedés.")
    add_bullet(doc, f"{len(overdue)} nyitott intézkedés belső céldátuma {AS_OF.isoformat()}-ig lejárt; ezekhez státusz- vagy újraütemezési döntés kell.")
    add_bullet(doc, "A 2026-06-26-i kézhezvételi dátum elfogadott baseline; külön címzetti kézbesítési igazolás nem áll rendelkezésre, ez nyilvántartott és elfogadott kockázat.")
    add_bullet(doc, "A kanonikus beadási határidő 2026-09-24. A beadást csak jogosult ember végezheti el.")

    doc.add_heading("2. Hatókör, alapok és korlátok", level=1)
    add_bullet(doc, "Kanonikus auditforrás: SRC-008, a hatóság felé beadott, aláírt Audit_Cert dokumentum; SRC-002 történeti archívum.")
    add_bullet(doc, "Minden AI-kimenet javaslat: nem jóváhagyás, nem bizonyíték és nem helyettesíti a szakmai, jogi vagy vezetői döntést.")
    add_bullet(doc, "Az unverified_internal forrásból származó műszaki állítások read-only ellenőrzésig feltételesek.")
    add_bullet(doc, "Éles rendszer-, jogosultság-, hálózati vagy konfigurációváltozás csak külön G3 jóváhagyással végezhető.")
    add_bullet(doc, "Fizetős megoldás csak a költségkapu teljes bemeneti csomagja és G5 döntés után indítható.")

    doc.add_heading("3. Irányítás és szerepek", level=1)
    roles = [
        ("IBF / szakmai felügyelet", "Lángi Zoltán", "FÜGGŐ: alkalmasság és formális kijelölési evidencia"),
        ("Jogi reviewer", "Dr. Berta Brigitta", "G2 jogi review"),
        ("Projektvezető", "Kóczán Mónika", "Koordináció"),
        ("Technikai végrehajtó", "Kollár Csaba / Serversystem Kft.", "FÜGGŐ: belső metALCOM kontrollgazda"),
        ("HR-kontrollgazda", "Koncz Erika", "HR és személyi biztonság"),
        ("Fizikai védelem", "Német Péter", "Fizikai és környezeti kontrollok"),
        ("Minden akció végrehajtási felelőse", "Pásztor András", "D-020 szerint"),
        ("Minden akció jóváhagyója", "Lángi Zoltán", "D-020 szerint; összeférhetetlenségi korlátokkal"),
        ("Vezetői szponzor", "NINCS MEGNEVEZVE", "FÜGGŐ – beadás előtt kijelölendő"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, text in enumerate(("Szerep", "Név", "Állapot / megjegyzés")):
        set_cell_shading(table.rows[0].cells[idx], NAVY)
        r = table.rows[0].cells[idx].paragraphs[0].add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE)
        r.font.size = Pt(8.5)
    set_repeat_table_header(table.rows[0])
    for role, name, state in roles:
        row = table.add_row()
        set_row_cant_split(row)
        for cell in row.cells:
            set_cell_margins(cell)
        if "FÜGGŐ" in state or "NINCS" in name:
            set_cell_shading(row.cells[2], PALE_RED)
        for cell, text in zip(row.cells, (role, name, state)):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            r.font.size = Pt(8)
            if "FÜGGŐ" in text or "NINCS" in text:
                r.bold = True
                r.font.color.rgb = RGBColor.from_string(RED)

    doc.add_heading("4. Beadás előtti kötelező függő elemek", level=1)
    add_banner(doc, "A következő pontok teljesítése nélkül a terv státusza nem lehet FINAL / APPROVED.", PALE_RED, RED)
    blockers = [
        ("B-01", "A-004", "Finding-regiszter G1 mintavétel és név szerinti reviewer sign-off."),
        ("B-02", "A-005 / A-043–A-127", "A 85 új kontrollszintű intézkedés és mapping G1 owner review-ja; a hiányzó kontrollgazdák kijelölése."),
        ("B-03", "A-036", "Aláírt RACI, formális IBF-kijelölés, vezetői szponzor és belső infrastruktúra-/incidenskontroll-gazda."),
        ("B-04", "Mind a 127 akció", "A D-035 határidő-ütemezés teljes tervre vonatkozó G4 ellenőrzése."),
        ("B-05", "A-022/A-024/A-026/A-027/A-028", "Az SRC-004 állítások read-only validációja vagy kifejezett feltételes minősítése."),
        ("B-06", "A-006/A-007", "G1 szakmai, G2 biztonsági-jogi és G4 vezetői beadási jóváhagyás."),
        ("B-07", "Végleges csomag", "Aláírt végleges fájl, verzió, védett SharePoint URI és SHA-256."),
        ("B-08", "Külső benyújtás", "Jogosult benyújtó, benyújtási visszaigazolás és címzetti/hatósági átvételi igazolás."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, text in enumerate(("Azonosító", "Kapcsolat", "Lezárási feltétel")):
        set_cell_shading(table.rows[0].cells[idx], RED)
        r = table.rows[0].cells[idx].paragraphs[0].add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE)
        r.font.size = Pt(8.5)
    set_repeat_table_header(table.rows[0])
    for blocker_id, reference, condition in blockers:
        row = table.add_row()
        set_row_cant_split(row)
        set_cell_shading(row.cells[0], PALE_RED)
        set_cell_shading(row.cells[2], PALE_RED)
        for cell in row.cells:
            set_cell_margins(cell)
        for cell, value in zip(row.cells, (blocker_id, reference, condition)):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            r.font.size = Pt(8)
            if cell in (row.cells[0], row.cells[2]):
                r.bold = True
                r.font.color.rgb = RGBColor.from_string(RED)

    doc.add_heading("5. Belső céldátumok ütemezési ellenőrzése", level=1)
    p = doc.add_paragraph(
        "A D-035 mind a 127 akcióhoz belső céldátumot rögzít. A 2027-09-30-i repeat-audit előtti 60. nap 2027-08-01; az operatív utolsó munkanap 2027-07-30."
    )
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, text in enumerate(("Akció", "Eredeti dátum", "Állapot", "Szükséges emberi lépés")):
        set_cell_shading(table.rows[0].cells[idx], RED)
        r = table.rows[0].cells[idx].paragraphs[0].add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE)
        r.font.size = Pt(8)
    set_repeat_table_header(table.rows[0])
    for action in overdue:
        row = table.add_row()
        set_row_cant_split(row)
        for cell in row.cells:
            set_cell_margins(cell)
        set_cell_shading(row.cells[1], PALE_RED)
        values = (action["action_id"], action["target_date"], action["status"], "Státusz igazolása vagy jóváhagyott újraütemezés")
        for cell, value in zip(row.cells, values):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            r.font.size = Pt(7.7)
            if cell is row.cells[1]:
                r.bold = True
                r.font.color.rgb = RGBColor.from_string(RED)

    doc.add_heading("6. Végrehajtási kapuk", level=1)
    gates = [
        ("G1", "Szakmai / kontrollgazdai review", "FÜGGŐ", "Findingok, mapping, hatókör, szakmai helyesség."),
        ("G2", "Biztonsági és jogi review", "FÜGGŐ", "Határidők, adatkezelés, személyes adatok, jogi megfelelőség."),
        ("G3", "Éles végrehajtási engedély", "Akciónként", "Bármely éles rendszer- vagy konfigurációváltozás előtt."),
        ("G4", "Vezetői / külső beadási jóváhagyás", "FÜGGŐ", "Aláírt végleges csomag, csatorna és jogosult benyújtó."),
        ("G5", "Költség- és beszerzési döntés", "Akciónként", "Fizetős megoldás előtt teljes költségkapu-csomaggal."),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, text in enumerate(("Kapu", "Jelentés", "Állapot", "Minimum eredmény")):
        set_cell_shading(table.rows[0].cells[idx], NAVY)
        r = table.rows[0].cells[idx].paragraphs[0].add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE)
        r.font.size = Pt(8.5)
    for gate, meaning, status, result in gates:
        row = table.add_row()
        set_row_cant_split(row)
        for cell in row.cells:
            set_cell_margins(cell)
        if status == "FÜGGŐ":
            set_cell_shading(row.cells[2], PALE_GOLD)
        for cell, value in zip(row.cells, (gate, meaning, status, result)):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            r.font.size = Pt(8)
            if value == "FÜGGŐ":
                r.bold = True
                r.font.color.rgb = RGBColor.from_string(GOLD)

    # Landscape action register.
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, landscape=True)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    add_header_footer(section, compact=True)
    doc.add_heading("7. Részletes intézkedési terv", level=1)
    p = doc.add_paragraph(
        f"Az alábbi mátrix a data/actions.csv kanonikus nyilvántartás teljes, {action_count} tételes tartalmát foglalja össze. "
        "A színezés az állapotot és a beadás előtti emberi teendőket mutatja; a dátumokat a D-035 jóváhagyott ütemezési szabálya rögzíti."
    )
    p.paragraph_format.space_after = Pt(8)
    add_family_coverage(doc, actions)
    doc.add_heading("7.2 Intézkedési törzsnyilvántartás", level=2)
    add_action_matrix(doc, actions)

    # Return to portrait for final controls and signatures.
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, landscape=False)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    add_header_footer(section)
    doc.add_heading("8. Beadás előtti ellenőrzőlista", level=1)
    checklist = [
        f"A {action_count} intézkedés szakmai tartalma, felelőse, határideje, eredménye és evidenciája G1 review-n megfelelt.",
        "A D-035 szerinti 127 belső céldátum és a 60 napos repeat-audit puffer ellenőrzése dokumentált.",
        "Az SRC-004 feltételes állítások validáltak vagy egyértelműen feltételesek maradtak.",
        "Az aláírt RACI, formális szerepkijelölések, vezetői szponzor és belső technikai kontrollgazda rendelkezésre áll.",
        "A jogi, adatvédelmi és információbiztonsági G2 review döntési rekordja rendelkezésre áll.",
        "A vezetői G4 jóváhagyás, a benyújtási csatorna és a jogosult benyújtó rögzített.",
        "A végleges DOCX/PDF változat védett SharePoint URI-ja és SHA-256 értéke rögzített.",
        "A kézi aláírással ellátott példány olvasható; a kézi aláírás teljes értékű elfogadását a D-031 baseline rögzíti.",
        "A benyújtási visszaigazolás és a címzetti/hatósági átvételi igazolás megőrzési helye kijelölt.",
    ]
    for item in checklist:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.3)
        p.paragraph_format.first_line_indent = Cm(-0.3)
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run("☐  ")
        r.font.name = "Segoe UI Symbol"
        r.font.size = Pt(11)
        p.add_run(item)

    doc.add_heading("9. Jóváhagyási és aláírási lap", level=1)
    add_banner(doc, "Csak a fenti ellenőrzőlista teljes lezárása után írható alá.", PALE_GOLD, GOLD)
    doc.add_paragraph()
    approvals = [
        ("G1 szakmai reviewer", "Név: ______________________________", "Dátum: ______________", "Aláírás: ______________________________"),
        ("G2 jogi reviewer", "Dr. Berta Brigitta", "Dátum: ______________", "Aláírás: ______________________________"),
        ("G2 IBF / biztonsági reviewer", "Lángi Zoltán", "Dátum: ______________", "Aláírás: ______________________________"),
        ("Projektvezető", "Kóczán Mónika", "Dátum: ______________", "Aláírás: ______________________________"),
        ("Végrehajtásért felelős", "Pásztor András", "Dátum: ______________", "Aláírás: ______________________________"),
        ("G4 jóváhagyó", "Lángi Zoltán", "Dátum: ______________", "Aláírás: ______________________________"),
        ("Jogosult benyújtó", "Név: ______________________________", "Dátum: ______________", "Aláírás: ______________________________"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for role, name, signed_at, signature in approvals:
        row = table.add_row()
        set_row_cant_split(row)
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        for cell in row.cells:
            set_cell_margins(cell, 130, 120, 130, 120)
        p = row.cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(role)
        r.bold = True
        r.font.size = Pt(9)
        p = row.cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(name)
        r.font.size = Pt(9)
        p = row.cells[1].add_paragraph(signed_at)
        p.paragraph_format.space_after = Pt(4)
        p.runs[0].font.size = Pt(8.5)
        p = row.cells[1].add_paragraph(signature)
        p.paragraph_format.space_after = Pt(0)
        p.runs[0].font.size = Pt(8.5)

    doc.add_heading("10. Végleges fájl és beadási evidencia", level=1)
    add_key_value_table(doc, [
        ("Végleges verzió", "FÜGGŐ – emberi jóváhagyás után kitöltendő"),
        ("Védett SharePoint URI", "FÜGGŐ – végleges, aláírt fájl feltöltése után kitöltendő"),
        ("SHA-256", "FÜGGŐ – a feltöltött, változatlan végleges fájlról számítandó"),
        ("Benyújtási azonosító", "FÜGGŐ – hatósági benyújtás után kitöltendő"),
        ("Átvételi igazolás URI", "FÜGGŐ – beérkezés után kitöltendő"),
    ])
    for table in doc.tables[-1:]:
        for row in table.rows:
            if "FÜGGŐ" in row.cells[1].text:
                set_cell_shading(row.cells[1], PALE_GOLD)
                for run in row.cells[1].paragraphs[0].runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(GOLD)

    core = doc.core_properties
    core.title = "metALCOM Zrt. – NIS2 helyreállítási cselekvési terv"
    core.subject = f"Beadásra előkészített, függő elemeket kiemelő {action_count} intézkedéses cselekvési terv"
    core.author = "langizoli"
    core.keywords = "NIS2, cselekvési terv, audit, metALCOM, G1, G2, G4"
    core.comments = "AI által előállított PROPOSAL; emberi jóváhagyás szükséges."
    doc.settings.update_fields_on_open = True
    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    print(build())
