"""Build one printable, signable decision sheet for every open DEF task."""

from __future__ import annotations

from datetime import datetime, timezone
import hashlib
import json
from pathlib import Path
import re
import sys
import tempfile
import zipfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from build_human_task_materials import (
    DOCUMENTS,
    LIGHT_BLUE,
    NAVY,
    OUTPUT_DIR,
    add_footer,
    build_document,
    configure_table,
    keep_table_row_together,
    repeat_table_header,
    set_cell_fill,
    style_run,
)

ROOT = Path(__file__).parents[1]
sys.path.insert(0, str(ROOT / "src"))

from nis2_harness.human_decision_documents import (
    build_specs as build_domain_specs,
    load_open_tasks as load_domain_open_tasks,
)


PACKAGE_PATH = ROOT / "data" / "human_execution_package.json"
MANIFEST_PATH = ROOT / "data" / "human_decision_document_manifest.json"
INDEX_PATH = OUTPUT_DIR / "NIS2_emberi_dontesi_dokumentumjegyzek.docx"
ARCHIVE_PATH = OUTPUT_DIR / "NIS2_emberi_dontesi_dokumentumcsomag.zip"
FIXED_ZIP_TIME = (2026, 7, 30, 12, 0, 0)

TITLES = {
    "DEF-008": "EIR- és exportforrás-kijelölés",
    "DEF-009": "Leltárfeltöltési és adatminőségi review",
    "DEF-010": "Agent QA és gold-case elfogadás",
    "DEF-011": "AI-használati szabályzat G2 jóváhagyása",
    "DEF-012": "Repeat-audit ütemterv G4/G5 döntése",
    "DEF-013": "Negyedéves riportütemezés jóváhagyása",
    "DEF-014": "Cselekvési terv G1/G2/G4 jóváhagyása",
    "DEF-015": "Belső portál architektúra G2/G3 döntése",
    "DEF-016": "Backup- és restore-teszt engedélyezése",
    "DEF-017": "Fizikai bejárás és adatkezelési jóváhagyás",
    "DEF-018": "Infrastruktúra health snapshot gyűjtési engedély",
    "DEF-019": "Licenc- és entitlement-review",
    "DEF-020": "Portál funkcionális scope és pilotdöntés",
    "DEF-021": "Naplózási és monitoring-baseline jóváhagyása",
    "DEF-022": "Karbantartási és változáskezelési baseline",
    "DEF-023": "Beszállítói kockázati baseline jóváhagyása",
    "DEF-024": "Exchange-függőségi gyűjtés és teszt engedély",
    "DEF-025": "Legacy megőrzési és restore-döntés",
    "DEF-026": "RDS-szeparációs assessment döntése",
    "DEF-027": "Technikai munkacsomagok kapudöntése",
    "DEF-028": "Működési kontrollcsomagok kapudöntése",
    "DEF-029": "Irányítási evidencialánc jóváhagyása",
    "DEF-030": "Szabályozási baseline-ok jóváhagyása",
    "DEF-031": "Folyamatos auditfelkészültségi agent pilot",
    "DEF-032": "Hitelesített portál- és Graph-pilot döntése",
    "DEF-033": "H-002 agent és Graph read pilot döntése",
    "DEF-034": "Átvett dokumentumok szakmai besorolása",
    "DEF-035": "SharePoint evidenciatár kontrollreview-ja",
    "DEF-036": "Kontrollkatalógus G1 és EIR-besorolási döntés",
    "DEF-037": "Lejárt akciók státusz- és határidő-egyeztetése",
    "DEF-038": "Többfelhasználós portálpilot indítási döntése",
}


def _slug(title: str) -> str:
    replacements = str.maketrans(
        "áéíóöőúüűÁÉÍÓÖŐÚÜŰ", "aeiooouuuAEIOOOUUU"
    )
    value = title.translate(replacements).lower()
    return re.sub(r"[^a-z0-9]+", "_", value).strip("_")[:70]


def _requirement_checks(text: str) -> list[str]:
    parts = [part.strip(" .") for part in text.split(";") if part.strip()]
    checks = parts[:10]
    if len(parts) > 10:
        checks.append("A további részletes követelmények a feladatrekord szerint ellenőrizve")
    checks.extend(
        [
            "A tényleges döntés vagy elkészült dokumentum csatolva",
            "A védett NIS2 SharePoint-hivatkozás és SHA-256 rögzítve",
            "A név szerinti reviewer és az időzónás review-idő rögzítve",
        ]
    )
    return checks


def load_open_tasks() -> list[dict[str, object]]:
    data = json.loads(PACKAGE_PATH.read_text(encoding="utf-8"))
    tasks: list[dict[str, object]] = []
    for wave in data["waves"]:
        for task in wave["tasks"]:
            if task["status"] != "OPEN_DEFERRED":
                continue
            tasks.append({**task, "wave": wave["wave_id"]})
    tasks.sort(key=lambda item: str(item["task_id"]))
    return tasks


def build_specs(tasks: list[dict[str, object]]) -> list[dict[str, object]]:
    curated = {str(item["task_id"]): dict(item) for item in DOCUMENTS}
    specs: list[dict[str, object]] = []
    for task in tasks:
        task_id = str(task["task_id"])
        if task_id in curated:
            spec = curated[task_id]
            spec.update(
                {
                    "wave": task["wave"],
                    "current_process_state": task["current_process_state"],
                    "must_be_completed_before": task["must_be_completed_before"],
                }
            )
        else:
            title = TITLES[task_id]
            spec = {
                "task_id": task_id,
                "filename": f"{task_id}_{_slug(title)}.docx",
                "title": title,
                "purpose": (
                    "A feladathoz szükséges emberi vizsgálat, kapudöntés, "
                    "mellékletek és jóváhagyási nyom egységes rögzítése. "
                    f"Elkészítendő eredmény: {task['required_result']}"
                ),
                "owner": task["owner"],
                "reviewer": task["approver"],
                "source_refs": task["related"],
                "wave": task["wave"],
                "current_process_state": task["current_process_state"],
                "must_be_completed_before": task["must_be_completed_before"],
                "steps": [
                    "Tekintse át a hivatkozott forrásokat és a jelenlegi folyamatállapotot.",
                    "Végezze el vagy igazolja a lapon felsorolt tényleges szakmai, jogi, biztonsági vagy műszaki ellenőrzéseket.",
                    "Csatolja a szükséges dokumentumokat, teszteredményeket és döntési mellékleteket; rögzítse a védett URI-kat és SHA-256 értékeket.",
                    "Jelölje a döntést, írja le a feltételeket vagy hiányokat, majd a jogosult reviewer írja alá a lapot.",
                    "A végleges, aláírt példányt töltse fel a védett NIS2 SharePoint-evidenciatárba.",
                ],
                "checks": _requirement_checks(str(task["required_result"])),
            }
        specs.append(spec)
    return specs


def normalize_docx(path: Path) -> None:
    """Normalize ZIP entry timestamps/order for deterministic DOCX output."""
    with tempfile.TemporaryDirectory() as temporary:
        normalized = Path(temporary) / path.name
        with zipfile.ZipFile(path, "r") as source:
            with zipfile.ZipFile(
                normalized, "w", compression=zipfile.ZIP_DEFLATED
            ) as target:
                for name in sorted(source.namelist()):
                    original = source.getinfo(name)
                    info = zipfile.ZipInfo(name, FIXED_ZIP_TIME)
                    info.compress_type = zipfile.ZIP_DEFLATED
                    info.external_attr = original.external_attr
                    info.create_system = original.create_system
                    target.writestr(info, source.read(name))
        path.write_bytes(normalized.read_bytes())


def build_index(
    specs: list[dict[str, object]], hashes: dict[str, str]
) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.right_margin = Pt(72)
    section.bottom_margin = section.left_margin = Pt(72)
    add_footer(section, "INDEX")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_run(
        title.add_run("NIS2 emberi döntési dokumentumjegyzék"),
        size=23,
        bold=True,
        color=NAVY,
    )
    doc.add_paragraph(
        "A jegyzék a 2026-07-30-i baseline 36 OPEN_DEFERRED tételéhez "
        "tartozó, nyomtatható és aláírható munkalapokat sorolja fel. Az "
        "aláírás csak a tényleges vizsgálattal és mellékletekkel együtt érvényes."
    )
    table = doc.add_table(rows=len(specs) + 1, cols=4)
    table.style = "Table Grid"
    configure_table(table, [900, 3000, 2700, 2760])
    repeat_table_header(table.rows[0])
    for row in table.rows:
        keep_table_row_together(row)
    for index, heading in enumerate(("ID", "Dokumentum", "Felelős", "Reviewer")):
        set_cell_fill(table.cell(0, index), LIGHT_BLUE)
        style_run(table.cell(0, index).paragraphs[0].add_run(heading), bold=True, color=NAVY)
    for row_index, spec in enumerate(specs, 1):
        values = (
            str(spec["task_id"]),
            str(spec["title"]),
            str(spec["owner"]),
            str(spec["reviewer"]),
        )
        for column, value in enumerate(values):
            style_run(table.cell(row_index, column).paragraphs[0].add_run(value), size=9)
    doc.add_heading("Fájlazonosítók", level=1)
    for spec in specs:
        filename = str(spec["filename"])
        paragraph = doc.add_paragraph()
        style_run(paragraph.add_run(f"{spec['task_id']}: "), bold=True)
        style_run(paragraph.add_run(f"{filename} · SHA-256: {hashes[filename]}"), size=8)
    fixed = datetime(2026, 8, 19, 12, 0, tzinfo=timezone.utc)
    doc.core_properties.created = fixed
    doc.core_properties.modified = fixed
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.title = "NIS2 emberi döntési dokumentumjegyzék"
    doc.save(INDEX_PATH)
    normalize_docx(INDEX_PATH)
    return INDEX_PATH


def main() -> None:
    tasks = load_domain_open_tasks(PACKAGE_PATH)
    specs = build_domain_specs(tasks, DOCUMENTS)
    if len(tasks) != 37 or len(specs) != 37:
        raise ValueError("Pontosan 37 OPEN_DEFERRED dokumentum szükséges.")
    outputs: list[Path] = []
    for spec in specs:
        output = build_document(spec)
        normalize_docx(output)
        outputs.append(output)
    hashes = {
        output.name: hashlib.sha256(output.read_bytes()).hexdigest()
        for output in outputs
    }
    index = build_index(specs, hashes)
    manifest = {
        "schema_version": 1,
        "status": "PRINTABLE_DRAFTS_READY_FOR_HUMAN_EXECUTION",
        "as_of": "2026-08-19",
        "formal_effect": False,
        "source_ref": "data/human_execution_package.json",
        "source_sha256": hashlib.sha256(PACKAGE_PATH.read_bytes()).hexdigest(),
        "document_count": len(outputs),
        "index": {
            "path": index.relative_to(ROOT).as_posix(),
            "sha256": hashlib.sha256(index.read_bytes()).hexdigest(),
        },
        "documents": [
            {
                "task_id": str(spec["task_id"]),
                "title": str(spec["title"]),
                "path": output.relative_to(ROOT).as_posix(),
                "sha256": hashes[output.name],
                "status": "DRAFT_REQUIRES_HUMAN_EXECUTION_AND_SIGNATURE",
            }
            for spec, output in zip(specs, outputs)
        ],
        "completion_warning": (
            "A nyomtatás és aláírás önmagában nem pótolja a feladatlapon "
            "előírt tényleges vizsgálatot, mellékletet vagy evidenciát."
        ),
    }
    MANIFEST_PATH.write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    with zipfile.ZipFile(
        ARCHIVE_PATH, "w", compression=zipfile.ZIP_DEFLATED
    ) as archive:
        for path in sorted([*outputs, index, MANIFEST_PATH]):
            info = zipfile.ZipInfo(
                path.relative_to(ROOT).as_posix(), FIXED_ZIP_TIME
            )
            info.compress_type = zipfile.ZIP_DEFLATED
            archive.writestr(info, path.read_bytes())
    print(
        f"{len(outputs)} munkalap, 1 jegyzék és 1 ZIP-csomag elkészült."
    )


if __name__ == "__main__":
    main()
