from __future__ import annotations

import csv
import importlib.util
import unittest
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SCRIPT = ROOT / "scripts" / "build_submission_action_plan.py"
DOCX = ROOT / "submission_materials" / "NIS2_cselekvesi_terv_beadasra_elokeszitett_2026-08-19.docx"


def load_builder():
    spec = importlib.util.spec_from_file_location("submission_action_plan", SCRIPT)
    if spec is None or spec.loader is None:
        raise RuntimeError("A cselekvésiterv-builder nem tölthető be.")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


class SubmissionActionPlanTest(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        cls.builder = load_builder()
        with (ROOT / "data" / "actions.csv").open(encoding="utf-8-sig", newline="") as handle:
            cls.actions = list(csv.DictReader(handle))

    def test_canonical_register_has_127_unique_actions_and_19_families(self) -> None:
        ids = [row["action_id"] for row in self.actions]
        families = {
            int(family)
            for row in self.actions
            for family in row["requirement_family"].split(";")
        }
        self.assertEqual(127, len(ids))
        self.assertEqual(127, len(set(ids)))
        self.assertEqual(set(range(1, 20)), families)

    def test_submission_blockers_and_pending_dates_are_explicit(self) -> None:
        self.assertEqual({"A-004", "A-005", "A-006", "A-036"}, set(self.builder.SUBMISSION_BLOCKERS))
        expected = {
            "A-008", "A-022", "A-023", "A-024", "A-025", "A-026", "A-027", "A-028", "A-042"
        } | {f"A-{index:03d}" for index in range(43, 128)}
        self.assertEqual(expected, self.builder.FIXED_DATE_REQUIRED)

    def test_generated_document_contains_all_actions_and_required_warnings(self) -> None:
        self.assertTrue(DOCX.exists())
        document = Document(DOCX)
        chunks = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                chunks.extend(cell.text for cell in row.cells)
        text = "\n".join(chunks)
        for index in range(1, 128):
            self.assertIn(f"A-{index:03d}", text)
        self.assertIn("PROPOSAL", text)
        self.assertIn("2026-09-24", text)
        self.assertIn("FÜGGŐ", text)
        self.assertIn("BEADÁSI BLOKKOLÓ", text)
        self.assertIn("NINCS MEGNEVEZVE", text)


if __name__ == "__main__":
    unittest.main()
